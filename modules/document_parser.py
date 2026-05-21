# File: authentilearn/modules/document_parser.py
import os
import io
import PyPDF2
import docx

def extract_text(uploaded_file) -> str:
    """
    Extracts text from a variety of file formats: PDF, DOCX, TXT, PY, JS, MD, HTML.
    """
    if uploaded_file is None:
        return ""
        
    filename = uploaded_file.name
    ext = os.path.splitext(filename)[1].lower()
    
    # Reset stream pointer
    uploaded_file.seek(0)
    
    if ext == ".pdf":
        try:
            reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text.strip()
        except Exception as e:
            return f"Error parsing PDF: {str(e)}"
            
    elif ext == ".docx":
        try:
            doc = docx.Document(uploaded_file)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            return "\n".join(text).strip()
        except Exception as e:
            return f"Error parsing DOCX: {str(e)}"
            
    elif ext in [".txt", ".py", ".js", ".md", ".html", ".css", ".json", ".java", ".c", ".cpp"]:
        try:
            content = uploaded_file.read()
            # Try utf-8
            try:
                return content.decode("utf-8")
            except UnicodeDecodeError:
                # Fallback to latin-1
                return content.decode("latin-1")
        except Exception as e:
            return f"Error parsing text file: {str(e)}"
            
    else:
        # Fallback to reading as text
        try:
            content = uploaded_file.read()
            return content.decode("utf-8", errors="ignore")
        except Exception as e:
            return f"Unsupported file type parsing error: {str(e)}"

def get_file_metadata(uploaded_file) -> dict:
    """
    Returns metadata dict with: name, size_kb, type, char_count, estimated_words
    """
    if uploaded_file is None:
        return {
            "name": "N/A",
            "size_kb": 0.0,
            "type": "Unknown",
            "char_count": 0,
            "estimated_words": 0
        }
        
    # Get file size
    uploaded_file.seek(0, os.SEEK_END)
    size_bytes = uploaded_file.tell()
    uploaded_file.seek(0) # reset
    
    filename = uploaded_file.name
    ext = os.path.splitext(filename)[1].lower()
    
    # Try to extract text to calculate word and character counts
    text = extract_text(uploaded_file)
    char_count = len(text)
    estimated_words = len(text.split())
    
    # Map extensions to nice readable types
    type_mapping = {
        ".pdf": "PDF Document",
        ".docx": "Word Document",
        ".txt": "Plain Text",
        ".py": "Python Script",
        ".js": "JavaScript File",
        ".md": "Markdown Document",
        ".html": "HTML Document",
        ".css": "CSS Stylesheet",
        ".json": "JSON File"
    }
    
    file_type = type_mapping.get(ext, f"{ext.upper()[1:]} File" if len(ext) > 1 else "Unknown")
    
    return {
        "name": filename,
        "size_kb": round(size_bytes / 1024.0, 2),
        "type": file_type,
        "char_count": char_count,
        "estimated_words": estimated_words
    }
